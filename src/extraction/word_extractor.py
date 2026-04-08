from docx import Document
from datetime import datetime
import logging

def extract_word(file_path):
    results = []

    try:
        doc = Document(file_path)
        logging.info(f"Opened Word file: {file_path}")

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        tables = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)
            tables.append(table_data)

        results.append({
            "content": "\n".join(paragraphs),
            "tables": tables,
            "metadata": {
                "file_name": file_path,
                "document_type": "word",
                "extraction_date": datetime.utcnow()
            }
        })

    except Exception as e:
        logging.error(f"Error processing Word file {file_path}: {e}")

    return results